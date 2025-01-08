import pandas as pd
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement, ns
from flask import Flask, request, jsonify, send_file, render_template
import os
import json

app = Flask(__name__)
UPLOAD_FOLDER = "uploads"
os.makedirs(UPLOAD_FOLDER, exist_ok=True)

class TranscriptProcessor:
    def __init__(self, input_vtt_path, options_path, save_intermediates=False):
        self.input_vtt_path = input_vtt_path
        self.options_path = options_path
        self.save_intermediates = save_intermediates
        self.options = self.load_options()
        self.vtt2docx_path = input_vtt_path.replace(".vtt", ".docx")
        self.step1_path = self.vtt2docx_path.replace(".docx", "_step1.docx")
        self.formatted_path = self.step1_path.replace("_step1.docx", "_formatted.docx")

    def load_options(self):
        # Load JSON configuration file
        with open(self.options_path, 'r') as f:
            return json.load(f)

    def apply_formatting(self, doc):
        for section in doc.sections:
            section.top_margin = Inches(self.options["margins"]["top"])
            section.bottom_margin = Inches(self.options["margins"]["bottom"])
            section.left_margin = Inches(self.options["margins"]["left"])
            section.right_margin = Inches(self.options["margins"]["right"])

            # Header
            header = section.header
            for paragraph in header.paragraphs:
                p_element = paragraph._element
                p_element.getparent().remove(p_element)

            for line in self.options["header"]["lines"]:
                header_paragraph = header.add_paragraph(line)
                header_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                header_paragraph.style.font.size = Pt(10)
                header_paragraph.paragraph_format.space_after = Pt(0)
                header_paragraph.paragraph_format.space_before = Pt(0)

            # Footer
            footer = section.footer
            for paragraph in footer.paragraphs:
                p_element = paragraph._element
                p_element.getparent().remove(p_element)

            footer_paragraph = footer.add_paragraph()
            footer_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            field_code = OxmlElement("w:fldSimple")
            field_code.set(ns.qn("w:instr"), "PAGE")
            footer_paragraph._p.append(field_code)

    def format_table(self, table):
        column_widths = self.options["table"]["column_widths"]
        alignments = self.options["table"]["alignments"]

        for col_idx, width in enumerate(column_widths):
            for cell in table.columns[col_idx].cells:
                cell.width = Inches(width)
                alignment = alignments[col_idx]
                if alignment == "left":
                    cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                elif alignment == "center":
                    cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                elif alignment == "right":
                    cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
                elif alignment == "justify":
                    cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    def vtt_to_docx(self):
        doc = Document()
        with open(self.input_vtt_path, "r", encoding="utf-8") as vtt_file:
            for line in vtt_file:
                doc.add_paragraph(line.strip())
        self.apply_formatting(doc)
        if self.save_intermediates:
            doc.save(self.vtt2docx_path)

    def transcript_to_word(self, input_path):
        doc = Document(input_path)
        line_index, timestamp, speaker, line = [], [], [], []
        current_index, current_timestamp, current_speaker, current_line = None, None, None, []

        for para in doc.paragraphs:
            text = para.text.strip()
            if text.isdigit():
                if current_index is not None:
                    line.append(" ".join(current_line))
                current_index = int(text)
                line_index.append(current_index)
                current_line = []
            elif "-->" in text:
                current_timestamp = text
                timestamp.append(current_timestamp)
            elif ":" in text:
                speaker_part, line_part = text.split(":", 1)
                current_speaker = speaker_part.strip().split()[0]
                current_line.append(line_part.strip())
                speaker.append(f"{current_speaker}:")
            else:
                current_line.append(text)

        if current_index is not None:
            line.append(" ".join(current_line))

        df = pd.DataFrame(
            {"Line Index": line_index, "Timestamp": timestamp, "Speaker": speaker, "Line": line}
        )
        output_doc = Document()
        table = output_doc.add_table(rows=0, cols=3)
        for _, row in df.iterrows():
            row_cells = table.add_row().cells
            row_cells[0].text = row["Timestamp"]
            row_cells[1].text = row["Speaker"]
            row_cells[2].text = row["Line"]

        self.format_table(table)
        self.apply_formatting(output_doc)
        if self.save_intermediates:
            output_doc.save(self.step1_path)

    def merge_speakers_and_clean_timestamps(self, input_path):
        def clean_timestamp(ts):
            return ts.split(" --> ")[0].split(".")[0] if "-->" in ts else ts

        doc = Document(input_path)
        table = doc.tables[0]
        data = [[cell.text.strip() for cell in row.cells] for row in table.rows]
        df = pd.DataFrame(data, columns=["Timestamp", "Speaker", "Line"])
        df["Timestamp"] = df["Timestamp"].apply(clean_timestamp)
        df["Speaker"] = df["Speaker"].apply(lambda x: x.split()[0])
        df_clean = df[df["Speaker"] != "R"]
        df_clean["group"] = (df_clean["Speaker"] != df_clean["Speaker"].shift()).cumsum()
        df_grouped = (
            df_clean.groupby("group")
            .agg({"Timestamp": "first", "Speaker": "first", "Line": " ".join})
            .reset_index(drop=True)
        )
        df_grouped["Speaker"] = df_grouped["Speaker"]

        output_doc = Document()
        table = output_doc.add_table(rows=0, cols=3)
        for _, row in df_grouped.iterrows():
            row_cells = table.add_row().cells
            row_cells[0].text = row["Timestamp"]
            row_cells[1].text = row["Speaker"]
            row_cells[2].text = row["Line"]

        self.format_table(table)
        self.apply_formatting(output_doc)
        output_doc.save(self.formatted_path)

    def process_transcript(self):
        self.vtt_to_docx()
        self.transcript_to_word(self.vtt2docx_path)
        self.merge_speakers_and_clean_timestamps(self.step1_path)
        return self.formatted_path

@app.route('/')
def home():
    return render_template('index.html')

@app.route('/process', methods=['POST'])
def process():
    if 'vtt_file' not in request.files:
        return jsonify({"error": "No file uploaded"}), 400

    vtt_file = request.files['vtt_file']
    if vtt_file.filename == '':
        return jsonify({"error": "No file selected"}), 400

    input_vtt_path = os.path.join(UPLOAD_FOLDER, vtt_file.filename)
    options_path = "options.json"
    vtt_file.save(input_vtt_path)

    processor = TranscriptProcessor(input_vtt_path, options_path, save_intermediates=False)
    output_path = processor.process_transcript()

    return send_file(output_path, as_attachment=True)

if __name__ == "__main__":
    app.run(debug=True, port=8000)
